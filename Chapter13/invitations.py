#! python3
# Invitations.py - make invitations for list from guests.txt

import docx, logging

logging.basicConfig(level=logging.DEBUG,
                    format='%(asctime)s - %(levelname)s - %(message)s')

# open guest.txt and create a list
guestText = open('Chapter13\\guests.txt', 'r')
guestList = guestText.read().splitlines()
logging.debug(guestList)

# TODO setting up word doc
doc = docx.Document()

for guest in guestList:
    # set up letter
    # It would be a pleasure to have the company of h2
    # <guest name> h1
    # at 11010 Memory Lane on the Evening of h2
    # April 1st h3
    # at 7 O'clock h2
    doc.add_paragraph(' ')
    doc.add_heading('It would be a pleasure to have the company of', 2)
    doc.add_heading(guest, 1)
    doc.add_heading('at 11011 Memory Lane on the Evening of', 2)
    doc.add_heading('April 1st', 3)
    doc.add_heading("at 7 O'clock", 4)
    doc.paragraphs[0].runs[0].add_break(docx.text.run.WD_BREAK.PAGE)

doc.save('Chapter13\\invitations.docx')


